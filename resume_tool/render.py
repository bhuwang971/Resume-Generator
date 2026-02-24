from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

from .schema import ResumeContent


class LayoutValidationError(Exception):
    pass


def _get(items: list[str], idx: int) -> str:
    if idx < len(items):
        return items[idx].strip()
    return ""


def build_context(resume: ResumeContent) -> dict[str, str]:
    context: dict[str, str] = {
        "summary": resume.professional_summary.strip(),
        "project2_name": resume.projects.project2.name.strip(),
        "project3_name": resume.projects.project3.name.strip(),
        "header_name": resume.header.name.strip(),
        "header_phone": resume.header.phone.strip(),
        "header_email": resume.header.email.strip(),
        "header_linkedin": resume.header.linkedin.strip(),
        "header_github": resume.header.github.strip(),
    }

    for idx in range(8):
        context[f"skills_{idx + 1}"] = _get(resume.technical_skills, idx)

    for idx in range(7):
        context[f"thorogood_{idx + 1}"] = _get(
            resume.work_experience.thorogood_bullets, idx
        )

    for idx in range(4):
        context[f"gta_{idx + 1}"] = _get(resume.work_experience.gwu_gta_bullets, idx)

    for idx in range(6):
        context[f"wtchtwr_{idx + 1}"] = _get(resume.projects.wtchtwr_bullets, idx)

    for idx in range(3):
        context[f"project2_{idx + 1}"] = _get(resume.projects.project2.bullets, idx)

    for idx in range(3):
        context[f"project3_{idx + 1}"] = _get(resume.projects.project3.bullets, idx)

    context["education_1"] = _get(resume.education_lines, 0)
    context["education_2"] = _get(resume.education_lines, 1)

    return context


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _normalize_text(text: str) -> str:
    return re.sub(r"[^a-z0-9]", "", (text or "").lower())


def _load_placeholder_indices(prepped_template_path: Path) -> dict[str, int]:
    doc = Document(str(prepped_template_path))
    indices: dict[str, int] = {}
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("{{") and text.endswith("}}"):
            indices[text] = idx
    return indices


def _required_indices(placeholders: dict[str, int], prefix: str, count: int) -> list[int]:
    indices: list[int] = []
    for idx in range(1, count + 1):
        key = f"{{{{{prefix}_{idx}}}}}"
        if key not in placeholders:
            raise LayoutValidationError(f"Missing placeholder in template: {key}")
        indices.append(placeholders[key])
    return indices


def _optional_indices(placeholders: dict[str, int], prefix: str, count: int) -> list[int]:
    indices: list[int] = []
    for idx in range(1, count + 1):
        key = f"{{{{{prefix}_{idx}}}}}"
        if key in placeholders:
            indices.append(placeholders[key])
    return indices


def _format_skill_line(paragraph: Paragraph, skill_line: str) -> None:
    line = skill_line.strip()
    label, sep, remainder = line.partition(":")

    if not sep:
        _replace_paragraph_text(paragraph, line)
        if paragraph.runs:
            paragraph.runs[0].bold = False
            paragraph.runs[0].underline = False
        return

    label_text = f"{label.strip()}:"
    remainder_text = remainder.strip()

    _replace_paragraph_text(paragraph, label_text)
    if paragraph.runs:
        paragraph.runs[0].bold = True
        paragraph.runs[0].underline = False

    if remainder_text:
        trailing = paragraph.add_run(f" {remainder_text}")
        trailing.bold = False
        trailing.underline = False


def _apply_skill_heading_format(
    output_path: Path, prepped_template_path: Path, resume: ResumeContent
) -> None:
    placeholders = _load_placeholder_indices(prepped_template_path)
    skill_indices = _required_indices(placeholders, "skills", 8)

    doc = Document(str(output_path))
    paragraphs = list(doc.paragraphs)

    for skill_idx, paragraph_idx in enumerate(skill_indices):
        if paragraph_idx < 0 or paragraph_idx >= len(paragraphs):
            raise LayoutValidationError(
                f"skills index {paragraph_idx} is out of range for rendered document."
            )
        _format_skill_line(paragraphs[paragraph_idx], _get(resume.technical_skills, skill_idx))

    # Remove duplicate skill lines between first skill and next placeholder block.
    start = min(skill_indices)
    next_indices = [idx for token, idx in placeholders.items() if token not in {f"{{{{skills_{i}}}}}" for i in range(1, 9)} and idx > start]
    end = min(next_indices) if next_indices else len(paragraphs)
    skill_norms = {_normalize_text(_get(resume.technical_skills, i)) for i in range(8)}
    skill_norms.discard("")

    for idx in range(start, min(end, len(paragraphs))):
        if idx in skill_indices:
            continue
        text = paragraphs[idx].text.strip()
        if not text:
            continue
        if _normalize_text(text) in skill_norms:
            _replace_paragraph_text(paragraphs[idx], "")

    doc.save(str(output_path))


def _is_bullet_word_paragraph(word_paragraph) -> bool:
    text = str(word_paragraph.Range.Text).replace("\r", "").strip()
    if not text:
        return False

    try:
        list_type = int(word_paragraph.Range.ListFormat.ListType)
    except Exception:  # noqa: BLE001
        list_type = 0

    if list_type != 0:
        return True
    return text.lstrip().startswith(("•", "-", "*"))


def _analyze_layout_with_word(
    docx_path: Path, bullet_indices: list[int]
) -> tuple[int, list[tuple[int, int, str]]]:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise LayoutValidationError(
            "Layout checks require Microsoft Word automation support. "
            "Install dependency with: pip install pywin32"
        ) from exc

    wd_statistic_lines = 1
    wd_statistic_pages = 2
    line_violations: list[tuple[int, int, str]] = []
    page_count = 0

    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(docx_path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        page_count = int(document.ComputeStatistics(wd_statistic_pages))

        max_paragraphs = int(document.Paragraphs.Count)
        for para_idx in sorted(set(bullet_indices)):
            word_para_idx = para_idx + 1
            if word_para_idx < 1 or word_para_idx > max_paragraphs:
                continue

            word_paragraph = document.Paragraphs(word_para_idx)
            if not _is_bullet_word_paragraph(word_paragraph):
                continue

            paragraph_range = word_paragraph.Range
            text = str(paragraph_range.Text).replace("\r", "").strip()
            if not text:
                continue

            line_count = int(paragraph_range.ComputeStatistics(wd_statistic_lines))
            if line_count > 1:
                line_violations.append((para_idx, line_count, text))
    except Exception as exc:  # noqa: BLE001
        raise LayoutValidationError(
            "Failed to run Word layout checks. Ensure Microsoft Word is installed."
        ) from exc
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()

    return page_count, line_violations


def _enforce_layout_constraints(output_path: Path, prepped_template_path: Path) -> None:
    placeholders = _load_placeholder_indices(prepped_template_path)
    bullet_indices = (
        _optional_indices(placeholders, "thorogood", 7)
        + _optional_indices(placeholders, "gta", 4)
        + _optional_indices(placeholders, "wtchtwr", 6)
        + _optional_indices(placeholders, "project2", 3)
        + _optional_indices(placeholders, "project3", 3)
    )

    page_count, bullet_violations = _analyze_layout_with_word(output_path, bullet_indices)

    errors: list[str] = []
    if page_count > 1:
        errors.append(f"Resume must be exactly 1 page, but generated {page_count} pages.")

    if bullet_violations:
        errors.append("Each bullet must fit in one line. Overflow bullets:")
        for para_idx, lines, text in bullet_violations:
            errors.append(f"- Paragraph {para_idx}: {lines} lines -> {text}")

    if errors:
        raise LayoutValidationError("\n".join(errors))


def render_resume_docx(
    prepped_template_path: str | Path, resume: ResumeContent, output_path: str | Path
) -> Path:
    prepped_path = Path(prepped_template_path)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    temp_output = output.with_name(f"{output.stem}.tmp{output.suffix}")

    template = DocxTemplate(str(prepped_path))
    template.render(build_context(resume))
    template.save(str(temp_output))

    _apply_skill_heading_format(temp_output, prepped_path, resume)
    _enforce_layout_constraints(temp_output, prepped_path)

    if output.exists():
        output.unlink()
    temp_output.replace(output)
    return output
