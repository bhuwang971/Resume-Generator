from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph


SECTION_HEADINGS = (
    "PROFESSIONAL SUMMARY",
    "TECHNICAL SKILLS",
    "WORK EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
)


class TemplatePrepError(Exception):
    pass


def _clean(text: str) -> str:
    return " ".join((text or "").strip().split())


def _is_bullet_paragraph(paragraph: Paragraph) -> bool:
    text = _clean(paragraph.text)
    if not text:
        return False

    style_name = ""
    if paragraph.style is not None and paragraph.style.name:
        style_name = paragraph.style.name.lower()

    if "list" in style_name or "bullet" in style_name or "number" in style_name:
        return True

    if text.startswith(("\u2022", "-", "*")):
        return True

    ppr = paragraph._p.pPr
    if ppr is not None and ppr.numPr is not None:
        return True

    return False


def _find_heading_index(paragraphs: list[Paragraph], heading: str) -> int:
    for idx, paragraph in enumerate(paragraphs):
        text = _clean(paragraph.text)
        if text == heading and text.upper() == text:
            return idx
    raise TemplatePrepError(f"Missing heading: {heading}")


def _get_non_empty_indices(
    paragraphs: list[Paragraph], start_idx: int, end_idx: int
) -> list[int]:
    indices: list[int] = []
    for idx in range(start_idx, end_idx):
        if _clean(paragraphs[idx].text):
            indices.append(idx)
    return indices


def empty_mapping_template() -> dict[str, Any]:
    return {
        "summary": 0,
        "skills": [0] * 8,
        "thorogood": [0] * 7,
        "gta": [0] * 4,
        "wtchtwr": [0] * 6,
        "project2_name": 0,
        "project2": [0] * 3,
        "project3_name": 0,
        "project3": [0] * 3,
    }


def _validate_int(value: Any, key: str) -> int:
    if not isinstance(value, int):
        raise TemplatePrepError(f"Mapping key '{key}' must be an integer index.")
    return value


def _validate_int_list(value: Any, key: str, expected_len: int) -> list[int]:
    if not isinstance(value, list):
        raise TemplatePrepError(f"Mapping key '{key}' must be a list of indices.")
    if len(value) != expected_len:
        raise TemplatePrepError(
            f"Mapping key '{key}' must have {expected_len} indices; found {len(value)}."
        )
    result: list[int] = []
    for item in value:
        if not isinstance(item, int):
            raise TemplatePrepError(f"Mapping key '{key}' contains non-integer value.")
        result.append(item)
    return result


def validate_mapping(mapping: dict[str, Any], paragraph_count: int) -> dict[str, Any]:
    required_keys = {
        "summary",
        "skills",
        "thorogood",
        "gta",
        "wtchtwr",
        "project2_name",
        "project2",
        "project3_name",
        "project3",
    }

    missing = required_keys - set(mapping.keys())
    if missing:
        missing_sorted = ", ".join(sorted(missing))
        raise TemplatePrepError(f"Mapping is missing keys: {missing_sorted}")

    normalized = {
        "summary": _validate_int(mapping["summary"], "summary"),
        "skills": _validate_int_list(mapping["skills"], "skills", 8),
        "thorogood": _validate_int_list(mapping["thorogood"], "thorogood", 7),
        "gta": _validate_int_list(mapping["gta"], "gta", 4),
        "wtchtwr": _validate_int_list(mapping["wtchtwr"], "wtchtwr", 6),
        "project2_name": _validate_int(mapping["project2_name"], "project2_name"),
        "project2": _validate_int_list(mapping["project2"], "project2", 3),
        "project3_name": _validate_int(mapping["project3_name"], "project3_name"),
        "project3": _validate_int_list(mapping["project3"], "project3", 3),
    }

    all_indices = [normalized["summary"], normalized["project2_name"], normalized["project3_name"]]
    all_indices.extend(normalized["skills"])
    all_indices.extend(normalized["thorogood"])
    all_indices.extend(normalized["gta"])
    all_indices.extend(normalized["wtchtwr"])
    all_indices.extend(normalized["project2"])
    all_indices.extend(normalized["project3"])

    for idx in all_indices:
        if idx < 0 or idx >= paragraph_count:
            raise TemplatePrepError(
                f"Paragraph index {idx} is out of range (0 to {paragraph_count - 1})."
            )

    return normalized


def auto_detect_mapping(template_path: str | Path) -> dict[str, Any]:
    doc = Document(str(template_path))
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        raise TemplatePrepError("Template has no paragraphs.")

    heading_idx = {heading: _find_heading_index(paragraphs, heading) for heading in SECTION_HEADINGS}

    summary_pool = _get_non_empty_indices(
        paragraphs,
        heading_idx["PROFESSIONAL SUMMARY"] + 1,
        heading_idx["TECHNICAL SKILLS"],
    )
    if not summary_pool:
        raise TemplatePrepError("Could not find summary paragraph after PROFESSIONAL SUMMARY.")

    skills_pool = _get_non_empty_indices(
        paragraphs,
        heading_idx["TECHNICAL SKILLS"] + 1,
        heading_idx["WORK EXPERIENCE"],
    )
    if len(skills_pool) < 8:
        raise TemplatePrepError(
            f"TECHNICAL SKILLS has {len(skills_pool)} non-empty lines; need at least 8."
        )

    work_pool = _get_non_empty_indices(
        paragraphs,
        heading_idx["WORK EXPERIENCE"] + 1,
        heading_idx["PROJECTS"],
    )
    work_bullets = [idx for idx in work_pool if _is_bullet_paragraph(paragraphs[idx])]
    if len(work_bullets) < 11:
        raise TemplatePrepError(
            "WORK EXPERIENCE must contain at least 11 detectable bullets (7 Thorogood + 4 GWU GTA). "
            f"Found {len(work_bullets)}."
        )

    projects_pool = _get_non_empty_indices(
        paragraphs,
        heading_idx["PROJECTS"] + 1,
        heading_idx["EDUCATION"],
    )
    project_bullets = [idx for idx in projects_pool if _is_bullet_paragraph(paragraphs[idx])]
    if len(project_bullets) < 12:
        raise TemplatePrepError(
            "PROJECTS must contain at least 12 detectable bullets (6 + 3 + 3). "
            f"Found {len(project_bullets)}."
        )

    project_name_lines = [
        idx
        for idx in projects_pool
        if _clean(paragraphs[idx].text) and not _is_bullet_paragraph(paragraphs[idx])
    ]
    if len(project_name_lines) < 3:
        raise TemplatePrepError(
            "Could not find at least 3 project title lines in PROJECTS section."
        )

    mapping = {
        "summary": summary_pool[0],
        "skills": skills_pool[:8],
        "thorogood": work_bullets[:7],
        "gta": work_bullets[7:11],
        "wtchtwr": project_bullets[:6],
        "project2_name": project_name_lines[1],
        "project2": project_bullets[6:9],
        "project3_name": project_name_lines[2],
        "project3": project_bullets[9:12],
    }

    return validate_mapping(mapping, len(paragraphs))


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _build_replacement_pairs(mapping: dict[str, Any]) -> list[tuple[int, str]]:
    pairs: list[tuple[int, str]] = []

    pairs.append((mapping["summary"], "{{summary}}"))

    for idx, paragraph_idx in enumerate(mapping["skills"], start=1):
        pairs.append((paragraph_idx, f"{{{{skills_{idx}}}}}"))

    for idx, paragraph_idx in enumerate(mapping["thorogood"], start=1):
        pairs.append((paragraph_idx, f"{{{{thorogood_{idx}}}}}"))

    for idx, paragraph_idx in enumerate(mapping["gta"], start=1):
        pairs.append((paragraph_idx, f"{{{{gta_{idx}}}}}"))

    for idx, paragraph_idx in enumerate(mapping["wtchtwr"], start=1):
        pairs.append((paragraph_idx, f"{{{{wtchtwr_{idx}}}}}"))

    pairs.append((mapping["project2_name"], "{{project2_name}}"))
    for idx, paragraph_idx in enumerate(mapping["project2"], start=1):
        pairs.append((paragraph_idx, f"{{{{project2_{idx}}}}}"))

    pairs.append((mapping["project3_name"], "{{project3_name}}"))
    for idx, paragraph_idx in enumerate(mapping["project3"], start=1):
        pairs.append((paragraph_idx, f"{{{{project3_{idx}}}}}"))

    return pairs


def prepare_template(
    template_path: str | Path,
    output_path: str | Path,
    mapping_output_path: str | Path | None = None,
    mapping: dict[str, Any] | None = None,
) -> dict[str, Any]:
    source_doc = Document(str(template_path))
    paragraphs = list(source_doc.paragraphs)

    if mapping is None:
        mapping = auto_detect_mapping(template_path)
    else:
        mapping = validate_mapping(mapping, len(paragraphs))

    for paragraph_idx, placeholder in _build_replacement_pairs(mapping):
        _replace_paragraph_text(paragraphs[paragraph_idx], placeholder)

    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    source_doc.save(str(output_file))

    if mapping_output_path is not None:
        mapping_file = Path(mapping_output_path)
        mapping_file.parent.mkdir(parents=True, exist_ok=True)
        mapping_file.write_text(json.dumps(mapping, indent=2), encoding="utf-8")

    return mapping


def list_paragraphs(template_path: str | Path) -> list[dict[str, Any]]:
    doc = Document(str(template_path))
    listing: list[dict[str, Any]] = []

    for idx, paragraph in enumerate(doc.paragraphs):
        style = paragraph.style.name if paragraph.style is not None else ""
        listing.append(
            {
                "index": idx,
                "text": _clean(paragraph.text),
                "style": style,
                "is_bullet": _is_bullet_paragraph(paragraph),
            }
        )

    return listing
